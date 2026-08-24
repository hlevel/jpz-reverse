#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档格式规范化脚本
将Word文档格式化为行业标准格式
版本：v1.0 (2026-05-21)
"""

import io
import sys
import os
import json

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 设置标准输出编码（解决Windows控制台编码问题）
import locale
try:
    # 尝试设置控制台编码
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    if hasattr(sys.stderr, 'reconfigure'):
        sys.stderr.reconfigure(encoding='utf-8')
except Exception:
    pass


# 格式标准配置
FORMAT_STANDARDS = {
    "government": {
        "name": "政府标书标准",
        "body": {
            "font": "仿宋_GB2312",
            "size": 14,
            "line_spacing": 28,
            "first_line_indent": 2,
            "alignment": "justify"
        },
        "headings": {
            "level1": {"font": "黑体", "size": 16, "bold": True, "alignment": "center"},
            "level2": {"font": "黑体", "size": 15, "bold": True, "alignment": "left"},
            "level3": {"font": "黑体", "size": 14, "bold": True, "alignment": "left"},
            "level4": {"font": "仿宋_GB2312", "size": 14, "bold": True, "alignment": "left"}
        },
        "table": {
            "header": {"font": "黑体", "size": 12, "bold": True, "alignment": "center"},
            "content": {"font": "仿宋_GB2312", "size": 10.5, "alignment": "center"}
        },
        "margins": {"top": 2.54, "bottom": 2.54, "left": 3.17, "right": 3.17}
    },
    "enterprise": {
        "name": "企业标书标准",
        "body": {
            "font": "宋体",
            "size": 12,
            "line_spacing": 18,
            "first_line_indent": 2,
            "alignment": "justify"
        },
        "headings": {
            "level1": {"font": "黑体", "size": 18, "bold": True, "alignment": "center"},
            "level2": {"font": "黑体", "size": 16, "bold": True, "alignment": "left"},
            "level3": {"font": "黑体", "size": 15, "bold": True, "alignment": "left"},
            "level4": {"font": "宋体", "size": 14, "bold": True, "alignment": "left"}
        },
        "table": {
            "header": {"font": "黑体", "size": 12, "bold": True, "alignment": "center"},
            "content": {"font": "宋体", "size": 10.5, "alignment": "center"}
        },
        "margins": {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54}
    },
    "highway": {
        "name": "高速公路标准",
        "body": {
            "font": "宋体",
            "size": 16,
            "line_spacing": 28,
            "first_line_indent": 2,
            "alignment": "justify"
        },
        "headings": {
            "level1": {"font": "黑体", "size": 18, "bold": True, "alignment": "center"},
            "level2": {"font": "黑体", "size": 16, "bold": True, "alignment": "left"},
            "level3": {"font": "黑体", "size": 15, "bold": True, "alignment": "left"},
            "level4": {"font": "宋体", "size": 14, "bold": True, "alignment": "left"}
        },
        "table": {
            "header": {"font": "黑体", "size": 12, "bold": True, "alignment": "center"},
            "content": {"font": "宋体", "size": 10.5, "alignment": "center"}
        },
        "margins": {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5}
    }
}


def set_run_font(run, font_name, font_size, bold=False):
    """设置run字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)

    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def format_heading(para, level, standard):
    """格式化标题"""
    heading_key = f"level{level}"
    heading_format = standard["headings"].get(heading_key, standard["headings"]["level1"])

    for run in para.runs:
        set_run_font(run, heading_format["font"], heading_format["size"], heading_format["bold"])

    if heading_format["alignment"] == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def format_body(para, standard):
    """格式化正文"""
    body_format = standard["body"]

    for run in para.runs:
        set_run_font(run, body_format["font"], body_format["size"])

    para.paragraph_format.line_spacing = Pt(body_format["line_spacing"])
    para.paragraph_format.first_line_indent = Pt(body_format["size"] * body_format["first_line_indent"])

    if body_format["alignment"] == "justify":
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def format_table(table, standard):
    """格式化表格"""
    table_format = standard["table"]

    # 格式化表头
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    set_run_font(run, table_format["header"]["font"],
                                table_format["header"]["size"],
                                table_format["header"]["bold"])

    # 格式化内容
    for row in table.rows[1:]:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    set_run_font(run, table_format["content"]["font"],
                                table_format["content"]["size"])


def format_docx(input_path, output_path=None, standard_name="government"):
    """格式化Word文档"""
    if not os.path.exists(input_path):
        print(f"[ERROR] 文件不存在: {input_path}")
        return False

    # 获取格式标准
    standard = FORMAT_STANDARDS.get(standard_name)
    if not standard:
        print(f"[ERROR] 未知的格式标准: {standard_name}")
        print(f"[INFO] 可用的标准: {', '.join(FORMAT_STANDARDS.keys())}")
        return False

    print(f"[INFO] 使用格式标准: {standard['name']}")

    # 打开文档
    doc = Document(input_path)

    # 设置页面边距
    margins = standard["margins"]
    for section in doc.sections:
        section.top_margin = Cm(margins["top"])
        section.bottom_margin = Cm(margins["bottom"])
        section.left_margin = Cm(margins["left"])
        section.right_margin = Cm(margins["right"])

    # 格式化段落
    for para in doc.paragraphs:
        # 检查是否是标题
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.replace('Heading ', '').replace('Heading', '1'))
            format_heading(para, level, standard)
        else:
            format_body(para, standard)

    # 格式化表格
    for table in doc.tables:
        format_table(table, standard)

    # 保存文档
    if output_path is None:
        output_path = input_path

    doc.save(output_path)
    print(f"[INFO] 文档已格式化: {output_path}")

    return True


def main():
    import argparse

    parser = argparse.ArgumentParser(description='Word文档格式化工具')
    parser.add_argument('input_file', help='输入Word文档路径')
    parser.add_argument('--output', '-o', help='输出文件路径')
    parser.add_argument('--standard', '-s', default='government',
                       choices=['government', 'enterprise', 'highway'],
                       help='格式标准（默认：government）')

    args = parser.parse_args()

    if not os.path.exists(args.input_file):
        print(f"[ERROR] 文件不存在: {args.input_file}")
        sys.exit(1)

    output_path = args.output or args.input_file

    success = format_docx(args.input_file, output_path, args.standard)
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
