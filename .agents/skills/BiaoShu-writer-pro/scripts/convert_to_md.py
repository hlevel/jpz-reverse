#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文档转Markdown脚本
将PDF/Word文档转换为Markdown格式，方便AI阅读和处理
版本：v1.0 (2026-05-21)

支持格式：
- .pdf PDF文件
- .doc 旧版Word文档（使用pywin32）
- .docx Word文档
- .txt 纯文本

输出格式：
- Markdown格式的文本文件
- 保留标题层级结构
- 表格转换为Markdown表格
- 图片位置标记
"""

import io
import os
import re
import sys
from pathlib import Path

# 设置标准输出编码（解决Windows控制台编码问题）
import locale
try:
    # 尝试设置控制台编码
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    if hasattr(sys.stderr, 'reconfigure'):
        sys.stderr.reconfigure(encoding='utf-8')
except Exception:
    pass


def convert_docx_to_md(file_path: str) -> str:
    """
    将Word文档(.docx)转换为Markdown格式

    Args:
        file_path: Word文档路径

    Returns:
        str: Markdown格式的文本
    """
    try:
        from docx import Document
    except ImportError:
        return "[错误] 需要安装 python-docx: pip install python-docx"

    try:
        doc = Document(file_path)
        md_lines = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                md_lines.append('')
                continue

            # 检测标题样式
            style_name = para.style.name.lower()

            if 'heading 1' in style_name or 'heading1' in style_name:
                md_lines.append(f'# {text}')
            elif 'heading 2' in style_name or 'heading2' in style_name:
                md_lines.append(f'## {text}')
            elif 'heading 3' in style_name or 'heading3' in style_name:
                md_lines.append(f'### {text}')
            elif 'heading 4' in style_name or 'heading4' in style_name:
                md_lines.append(f'#### {text}')
            elif 'heading 5' in style_name or 'heading5' in style_name:
                md_lines.append(f'##### {text}')
            elif 'list' in style_name:
                # 列表项
                md_lines.append(f'- {text}')
            else:
                # 普通段落
                md_lines.append(text)

        # 提取表格
        for table in doc.tables:
            md_lines.append('')  # 空行分隔

            # 表头
            if len(table.rows) > 0:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                md_lines.append('| ' + ' | '.join(headers) + ' |')
                md_lines.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')

                # 数据行
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    md_lines.append('| ' + ' | '.join(cells) + ' |')

            md_lines.append('')  # 空行分隔

        return '\n'.join(md_lines)

    except Exception as e:
        return f"[错误] DOCX转换失败: {str(e)}"


def convert_doc_to_md(file_path: str) -> str:
    """
    将旧版Word文档(.doc)转换为Markdown格式
    注意：此功能仅在Windows系统上可用

    Args:
        file_path: Word文档路径

    Returns:
        str: Markdown格式的文本
    """
    try:
        import win32com.client
        import pythoncom
    except ImportError:
        return "[错误] 需要安装 pywin32: pip install pywin32"

    try:
        # 初始化COM
        pythoncom.CoInitialize()

        # 创建Word应用程序对象
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        try:
            # 打开文档
            doc = word.Documents.Open(os.path.abspath(file_path))

            # 提取全部文本
            content = doc.Content.Text

            # 关闭文档
            doc.Close(False)

            # 简单处理为Markdown格式
            lines = content.split('\r')
            md_lines = []

            for line in lines:
                line = line.strip()
                if not line:
                    md_lines.append('')
                    continue

                # 简单的标题检测（基于数字编号）
                if re.match(r'^第[一二三四五六七八九十\d]+章', line):
                    md_lines.append(f'# {line}')
                elif re.match(r'^\d+\.\s+', line):
                    md_lines.append(f'## {line}')
                elif re.match(r'^\d+\.\d+\s+', line):
                    md_lines.append(f'### {line}')
                else:
                    md_lines.append(line)

            return '\n'.join(md_lines)

        finally:
            # 退出Word应用程序
            word.Quit()
            pythoncom.CoUninitialize()

    except Exception as e:
        return f"[错误] DOC转换失败: {str(e)}\n请确保已安装Microsoft Word或WPS Office"


def fix_encoding(text: str) -> str:
    """
    修复编码问题
    检测并修复UTF-8文本被错误用GBK/Latin-1解码的情况

    Args:
        text: 可能包含编码错误的文本

    Returns:
        str: 修复后的文本
    """
    if not text:
        return text

    # 检测是否包含典型的编码错误特征
    # UTF-8的中文被错误用GBK解码会出现这些字符
    garbled_chars = set('鎷鏍浜細璞戒俊鐢靛瓙绉戞妧闆嗗洟鏈夐檺鍏徃鎷涙爣浠ｇ悊鏈烘瀯锛氫腑寤哄北娌冲缓璁剧鐞嗛泦鍥㈡湁闄愬叕鍙')

    # 计算乱码字符比例
    garbled_count = sum(1 for c in text if c in garbled_chars)
    garbled_ratio = garbled_count / len(text) if len(text) > 0 else 0

    # 如果乱码字符比例超过10%，尝试修复
    if garbled_ratio > 0.1:
        try:
            # 尝试将文本编码为Latin-1，然后用UTF-8解码
            fixed = text.encode('latin-1').decode('utf-8')
            return fixed
        except (UnicodeDecodeError, UnicodeEncodeError):
            pass

        try:
            # 尝试将文本编码为GBK，然后用UTF-8解码
            fixed = text.encode('gbk').decode('utf-8')
            return fixed
        except (UnicodeDecodeError, UnicodeEncodeError):
            pass

    return text


def convert_pdf_to_md(file_path: str) -> str:
    """
    将PDF文件转换为Markdown格式

    Args:
        file_path: PDF文件路径

    Returns:
        str: Markdown格式的文本
    """
    try:
        import pdfplumber
    except ImportError:
        return "[错误] 需要安装 pdfplumber: pip install pdfplumber"

    try:
        md_lines = []

        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                # 添加分页标记
                if i > 0:
                    md_lines.append(f'\n---\n<!-- 第{i + 1}页 -->\n')

                # 提取文本
                text = page.extract_text()
                if text:
                    # 修复编码问题
                    text = fix_encoding(text)

                    # 处理文本，尝试识别标题
                    lines = text.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            md_lines.append('')
                            continue

                        # 简单的标题检测
                        if re.match(r'^第[一二三四五六七八九十\d]+章', line):
                            md_lines.append(f'# {line}')
                        elif re.match(r'^\d+\.\s+[A-Z\u4e00-\u9fa5]', line) and len(line) < 50:
                            md_lines.append(f'## {line}')
                        elif re.match(r'^\d+\.\d+\s+', line) and len(line) < 50:
                            md_lines.append(f'### {line}')
                        elif re.match(r'^\d+\.\d+\.\d+\s+', line) and len(line) < 50:
                            md_lines.append(f'#### {line}')
                        else:
                            md_lines.append(line)

                # 提取表格
                tables = page.extract_tables()
                for table in tables:
                    if table and len(table) > 0:
                        md_lines.append('')

                        # 处理表头
                        headers = table[0]
                        if headers:
                            headers = [str(cell).strip() if cell else '' for cell in headers]
                            # 修复表头编码
                            headers = [fix_encoding(h) for h in headers]
                            md_lines.append('| ' + ' | '.join(headers) + ' |')
                            md_lines.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')

                        # 处理数据行
                        for row in table[1:]:
                            if row:
                                cells = [str(cell).strip() if cell else '' for cell in row]
                                # 修复单元格编码
                                cells = [fix_encoding(c) for c in cells]
                                md_lines.append('| ' + ' | '.join(cells) + ' |')

                        md_lines.append('')

        return '\n'.join(md_lines)

    except Exception as e:
        return f"[错误] PDF转换失败: {str(e)}"


def convert_txt_to_md(file_path: str) -> str:
    """
    将纯文本文件转换为Markdown格式

    Args:
        file_path: 文本文件路径

    Returns:
        str: Markdown格式的文本
    """
    try:
        # 尝试不同编码
        content = None
        for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            return "[错误] 无法读取文件，编码识别失败"

        # 处理为Markdown格式
        lines = content.split('\n')
        md_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                md_lines.append('')
                continue

            # 简单的标题检测
            if re.match(r'^第[一二三四五六七八九十\d]+章', line):
                md_lines.append(f'# {line}')
            elif re.match(r'^\d+\.\s+[A-Z\u4e00-\u9fa5]', line) and len(line) < 50:
                md_lines.append(f'## {line}')
            elif re.match(r'^\d+\.\d+\s+', line) and len(line) < 50:
                md_lines.append(f'### {line}')
            else:
                md_lines.append(line)

        return '\n'.join(md_lines)

    except Exception as e:
        return f"[错误] TXT转换失败: {str(e)}"


def convert_to_markdown(input_file: str, output_file: str = None) -> str:
    """
    将文档转换为Markdown格式

    Args:
        input_file: 输入文件路径
        output_file: 输出文件路径（可选）

    Returns:
        str: Markdown格式的文本
    """
    if not os.path.exists(input_file):
        return f"[错误] 文件不存在: {input_file}"

    ext = os.path.splitext(input_file)[1].lower()

    # 选择转换函数
    converters = {
        '.txt': convert_txt_to_md,
        '.doc': convert_doc_to_md,
        '.docx': convert_docx_to_md,
        '.pdf': convert_pdf_to_md,
    }

    converter = converters.get(ext)
    if not converter:
        supported = ', '.join(converters.keys())
        return f"[错误] 不支持的文件格式: {ext}\n支持的格式: {supported}"

    print(f"[INFO] 正在转换: {input_file}")
    print(f"[INFO] 文件格式: {ext}")

    # 执行转换
    md_content = converter(input_file)

    # 检查是否转换成功
    if md_content.startswith("[错误]"):
        return md_content

    # 保存到文件
    if output_file:
        output_dir = os.path.dirname(output_file)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(md_content)

        print(f"[INFO] 已保存到: {output_file}")

    return md_content


def main():
    import argparse

    parser = argparse.ArgumentParser(description='文档转Markdown工具')
    parser.add_argument('input_file', help='输入文件路径')
    parser.add_argument('--output', '-o', help='输出文件路径（默认：输入文件名.md）')

    args = parser.parse_args()

    if not os.path.exists(args.input_file):
        print(f"[ERROR] 文件不存在: {args.input_file}")
        sys.exit(1)

    # 确定输出路径
    if args.output:
        output_file = args.output
    else:
        input_path = Path(args.input_file)
        output_file = str(input_path.parent / f"{input_path.stem}.md")

    print(f"[INFO] 开始转换...")
    print(f"[INFO] 输入文件: {args.input_file}")
    print(f"[INFO] 输出文件: {output_file}")

    # 执行转换
    md_content = convert_to_markdown(args.input_file, output_file)

    if md_content.startswith("[错误]"):
        print(f"\n{md_content}")
        sys.exit(1)
    else:
        # 统计信息
        lines = md_content.split('\n')
        chars = len(md_content)

        print(f"\n[INFO] 转换完成!")
        print(f"[INFO] 行数: {len(lines)}")
        print(f"[INFO] 字符数: {chars}")


if __name__ == '__main__':
    main()
