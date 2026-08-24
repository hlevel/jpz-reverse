#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
招标文件解析脚本 - 支持 txt, pdf, doc, docx, xlsx 格式
版本：v1.1 (2026-05-21)

功能：
1. 解析多种格式的招标文件
2. 提取文本内容和表格数据
3. 扫描版PDF检测和提示
4. 支持旧版Word文档(.doc)解析

支持格式：
- .txt 纯文本
- .doc 旧版Word文档（使用pywin32）
- .docx Word文档
- .pdf PDF文件（文本型/扫描版）
- .xlsx Excel表格
"""

import io
import sys
import os
from pathlib import Path

# 设置标准输出编码（解决Windows控制台编码问题）
import locale
try:
    # 尝试设置控制台编码
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    if hasattr(sys.stderr, 'reconfigure'):
        sys.stderr.reconfigure(encoding='utf-8')
except Exception:
    pass


def fix_encoding(text: str) -> str:
    """
    修复编码问题
    检测并修复UTF-8文本被错误用GBK/Latin-1解码的情况

    Args:
        text: 可能包含编码错误的文本

    Returns:
        str: 修复后的文本
    """
    if not text:
        return text

    # 检测是否包含典型的编码错误特征
    # UTF-8的中文被错误用GBK解码会出现这些字符
    garbled_chars = set('鎷鏍浜細璞戒俊鐢靛瓙绉戞妧闆嗗洟鏈夐檺鍏徃鎷涙爣浠ｇ悊鏈烘瀯锛氫腑寤哄北娌冲缓璁剧鐞嗛泦鍥缓璁剧鐞嗛泦鍥㈡湁闄愬叕鍙')

    # 计算乱码字符比例
    garbled_count = sum(1 for c in text if c in garbled_chars)
    garbled_ratio = garbled_count / len(text) if len(text) > 0 else 0

    # 如果乱码字符比例超过10%，尝试修复
    if garbled_ratio > 0.1:
        try:
            # 尝试将文本编码为Latin-1，然后用UTF-8解码
            fixed = text.encode('latin-1').decode('utf-8')
            return fixed
        except (UnicodeDecodeError, UnicodeEncodeError):
            pass

        try:
            # 尝试将文本编码为GBK，然后用UTF-8解码
            fixed = text.encode('gbk').decode('utf-8')
            return fixed
        except (UnicodeDecodeError, UnicodeEncodeError):
            pass

    return text


def parse_txt(file_path):
    """解析Txt文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='gbk') as f:
                return f.read()
        except Exception as e:
            return f"[错误] TXT读取失败: {str(e)}"


def parse_doc(file_path):
    """
    解析旧版Word文档(.doc) - 使用pywin32
    注意：此功能仅在Windows系统上可用
    """
    try:
        import win32com.client
        import pythoncom
    except ImportError:
        return "[错误] 需要安装 pywin32: pip install pywin32"

    try:
        # 初始化COM
        pythoncom.CoInitialize()

        # 创建Word应用程序对象
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        try:
            # 打开文档
            doc = word.Documents.Open(os.path.abspath(file_path))

            # 提取全部文本
            content = doc.Content.Text

            # 关闭文档
            doc.Close(False)

            return content
        finally:
            # 退出Word应用程序
            word.Quit()
            pythoncom.CoUninitialize()

    except Exception as e:
        return f"[错误] DOC解析失败: {str(e)}\n请确保已安装Microsoft Word或WPS Office"


def parse_docx(file_path):
    """解析Word文档(.docx)"""
    try:
        from docx import Document
    except ImportError:
        return "[错误] 需要安装 python-docx: pip install python-docx"

    try:
        doc = Document(file_path)
        paragraphs = []

        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    paragraphs.append(f'[表格] {row_text}')

        return '\n'.join(paragraphs)
    except Exception as e:
        return f"[错误] DOCX解析失败: {str(e)}"


def check_pdf_is_scanned(file_path):
    """
    检测PDF是否为扫描版
    返回: (is_scanned: bool, message: str)
    """
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            if len(pdf.pages) == 0:
                return True, "[警告] PDF页数为0，无法解析"

            # 尝试提取前3页文字
            first_page_text = ""
            for page in pdf.pages[:3]:
                text = page.extract_text()
                if text:
                    first_page_text = text
                    break

            if not first_page_text or len(first_page_text.strip()) < 50:
                return True, "[警告] PDF可能为扫描版（检测到文字层缺失或极少）"

            return False, "[信息] PDF为文本型，可正常解析"
    except Exception as e:
        return True, f"[错误] PDF解析异常: {str(e)}"


def parse_pdf(file_path):
    """
    解析PDF文件
    支持：文本型PDF直接提取
    提示：扫描版PDF需人工后处理
    """
    # 首先检测是否为扫描版
    is_scanned, check_msg = check_pdf_is_scanned(file_path)
    print(check_msg)

    if is_scanned:
        # 尝试备用方案
        try:
            import PyPDF2
            text_parts = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text and len(text.strip()) > 50:
                        text_parts.append(text)

            if text_parts:
                return "[提示] 使用PyPDF2提取到部分文字，可能不完整：\n\n" + '\n\n'.join(text_parts[:5])
            else:
                return """[错误] PDF为扫描版，无法自动提取文字

处理建议：
1. 使用ABBYY FineReader、Adobe Acrobat等OCR工具进行文字识别
2. 或手动复制PDF中的文字内容为txt文件
3. 识别后的文本可重新提交解析

推荐工具：
- Adobe Acrobat: 菜单 → 工具 → 扫描和OCR → 识别文本
- 在线OCR: https://www.ilovepdf.com/ocr
- ABBYY FineReader（专业级，精度最高）
"""
        except ImportError:
            return "[错误] 需要安装 PyPDF2: pip install PyPDF2"

    # 文本型PDF，使用pdfplumber
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    # 修复编码问题
                    text = fix_encoding(text)
                    text_parts.append(text)
        return '\n\n'.join(text_parts)
    except ImportError:
        return "[错误] 需要安装 pdfplumber: pip install pdfplumber"


def parse_xlsx(file_path):
    """解析Excel文件"""
    try:
        import openpyxl
    except ImportError:
        return "[错误] 需要安装 openpyxl: pip install openpyxl"

    try:
        text_parts = []
        wb = openpyxl.load_workbook(file_path, data_only=True)
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f'=== Sheet: {sheet_name} ===')
            for row in sheet.iter_rows(values_only=True):
                row_text = ' | '.join([str(cell) if cell is not None else '' for cell in row])
                if row_text.strip():
                    text_parts.append(row_text)
        return '\n'.join(text_parts)
    except Exception as e:
        return f"[错误] XLSX解析失败: {str(e)}"


def parse_file(file_path):
    """根据文件扩展名解析文件"""
    if not os.path.exists(file_path):
        return f"[错误] 文件不存在: {file_path}"

    ext = os.path.splitext(file_path)[1].lower()

    parsers = {
        '.txt': parse_txt,
        '.doc': parse_doc,
        '.docx': parse_docx,
        '.pdf': parse_pdf,
        '.xlsx': parse_xlsx,
        '.xls': parse_xlsx,
    }

    parser = parsers.get(ext)
    if not parser:
        supported = ', '.join(parsers.keys())
        return f"[错误] 不支持的文件格式: {ext}\n支持的格式: {supported}"

    try:
        return parser(file_path)
    except Exception as e:
        return f"[错误] 解析失败: {str(e)}\n请检查文件是否损坏或加密"


def main():
    if len(sys.argv) < 2:
        print('用法: python3 parse_bid_file.py <文件路径>')
        print('支持格式: txt, doc, docx, pdf, xlsx, xls')
        print('')
        print('PDF处理说明:')
        print('  - 文本型PDF: 直接提取文字')
        print('  - 扫描版PDF: 提示用户使用OCR工具')
        print('')
        print('DOC处理说明:')
        print('  - 旧版Word文档(.doc)需要安装pywin32: pip install pywin32')
        print('  - 需要系统安装Microsoft Word或WPS Office')
        sys.exit(1)

    file_path = sys.argv[1]
    print(f"正在解析: {file_path}")
    print("-" * 50)

    content = parse_file(file_path)
    print(content)


if __name__ == '__main__':
    main()
