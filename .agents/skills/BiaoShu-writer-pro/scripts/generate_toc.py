#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
目录生成脚本
生成标书目录页
版本：v1.0 (2026-05-21)
"""

import io
import os
import sys

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 设置标准输出编码（解决Windows控制台编码问题）
import locale
try:
    # 尝试设置控制台编码
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    if hasattr(sys.stderr, 'reconfigure'):
        sys.stderr.reconfigure(encoding='utf-8')
except Exception:
    pass


def set_run_font(run, font_name, font_size, bold=False):
    """设置run字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold

    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_toc_field(paragraph):
    """添加目录域代码"""
    # 添加域代码开始
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar)

    # 添加域代码内容
    run = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._element.append(instrText)

    # 添加域代码分隔
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar)

    # 添加占位文本
    run = paragraph.add_run("（请右键更新目录）")
    set_run_font(run, '宋体', 12)

    # 添加域代码结束
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar)


def generate_toc(chapters: list = None) -> Document:
    """
    生成目录页

    Args:
        chapters: 章节列表，格式为 [{"id": 1, "title": "章节标题", "sections": [...]}, ...]

    Returns:
        Document: Word文档对象
    """
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 添加目录标题
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("目  录")
    set_run_font(run, '黑体', 22, bold=True)

    # 添加空行
    doc.add_paragraph()

    # 添加目录域代码（自动目录）
    para = doc.add_paragraph()
    add_toc_field(para)

    # 如果提供了章节信息，添加静态目录
    if chapters:
        doc.add_paragraph()  # 空行分隔

        for chapter in chapters:
            chapter_id = chapter.get('id', '')
            chapter_title = chapter.get('title', '')

            # 添加章标题
            para = doc.add_paragraph()
            run = para.add_run(f"第{chapter_id}章  {chapter_title}")
            set_run_font(run, '黑体', 14, bold=True)

            # 添加节标题
            sections = chapter.get('sections', [])
            for i, section_title in enumerate(sections, 1):
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(1)
                run = para.add_run(f"{chapter_id}.{i}  {section_title}")
                set_run_font(run, '宋体', 12)

    return doc


def generate_static_toc(chapters: list, output_path: str):
    """生成静态目录"""
    doc = generate_toc(chapters)
    doc.save(output_path)
    print(f"[INFO] 目录已生成: {output_path}")


def generate_toc_from_outline(outline_file: str, output_path: str):
    """从大纲文件生成目录"""
    import json

    with open(outline_file, 'r', encoding='utf-8') as f:
        outline = json.load(f)

    chapters = outline.get('outline', [])
    generate_static_toc(chapters, output_path)


def main():
    import argparse

    parser = argparse.ArgumentParser(description='标书目录生成工具')
    parser.add_argument('--outline', '-i', help='大纲JSON文件路径')
    parser.add_argument('--output', '-o', default='toc.docx', help='输出文件路径')
    parser.add_argument('--auto', '-a', action='store_true', help='生成自动目录')

    args = parser.parse_args()

    print(f"[INFO] 开始生成目录...")

    if args.outline:
        generate_toc_from_outline(args.outline, args.output)
    else:
        doc = generate_toc()
        doc.save(args.output)
        print(f"[INFO] 目录已生成: {args.output}")


if __name__ == '__main__':
    main()
