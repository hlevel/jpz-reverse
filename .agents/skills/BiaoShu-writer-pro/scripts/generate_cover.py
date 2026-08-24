#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
封面生成脚本
生成标书封面页
版本：v1.0 (2026-05-21)
"""

import io
import os
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 设置标准输出编码（解决Windows控制台编码问题）
import locale
try:
    # 尝试设置控制台编码
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    if hasattr(sys.stderr, 'reconfigure'):
        sys.stderr.reconfigure(encoding='utf-8')
except Exception:
    pass


def set_run_font(run, font_name, font_size, bold=False, color=None):
    """设置run字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold

    if color:
        run.font.color.rgb = RGBColor(*color)

    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def add_empty_paragraph(doc, count=1):
    """添加空段落"""
    for _ in range(count):
        doc.add_paragraph()


def generate_cover(project_name: str, bidder_name: str = None, date: str = None) -> Document:
    """
    生成标书封面

    Args:
        project_name: 项目名称
        bidder_name: 投标单位名称
        date: 日期

    Returns:
        Document: Word文档对象
    """
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    # 添加空行
    add_empty_paragraph(doc, 3)

    # 添加LOGO占位符
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("[公司LOGO]")
    set_run_font(run, '黑体', 14, color=(128, 128, 128))

    # 添加空行
    add_empty_paragraph(doc, 2)

    # 添加项目名称
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(project_name)
    set_run_font(run, '华文中宋', 36, bold=True)

    # 添加空行
    add_empty_paragraph(doc, 1)

    # 添加"技术标"或"投标文件"
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("技术标")
    set_run_font(run, '华文中宋', 28, bold=True)

    # 添加空行
    add_empty_paragraph(doc, 4)

    # 添加投标单位
    if bidder_name:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"投标单位：{bidder_name}")
        set_run_font(run, '仿宋_GB2312', 16)

        add_empty_paragraph(doc, 1)

    # 添加日期
    if date is None:
        date = datetime.now().strftime("%Y年%m月%d日")

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(date)
    set_run_font(run, '仿宋_GB2312', 16)

    return doc


def generate_cover_from_info(info: dict, output_path: str):
    """从信息字典生成封面"""
    project_name = info.get('project_name', '投标项目')
    bidder_name = info.get('bidder_name', None)
    date = info.get('date', None)

    doc = generate_cover(project_name, bidder_name, date)
    doc.save(output_path)
    print(f"[INFO] 封面已生成: {output_path}")


def main():
    import argparse

    parser = argparse.ArgumentParser(description='标书封面生成工具')
    parser.add_argument('--project', '-p', required=True, help='项目名称')
    parser.add_argument('--bidder', '-b', help='投标单位名称')
    parser.add_argument('--date', '-d', help='日期（默认：当前日期）')
    parser.add_argument('--output', '-o', default='cover.docx', help='输出文件路径')

    args = parser.parse_args()

    print(f"[INFO] 开始生成封面...")

    doc = generate_cover(args.project, args.bidder, args.date)
    doc.save(args.output)

    print(f"[INFO] 封面已生成: {args.output}")


if __name__ == '__main__':
    main()
