#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown转Word文档脚本
将Markdown格式的标书内容转换为Word文档
版本：v1.0 (2026-05-21)
"""

import io
import sys
import os
import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 设置标准输出编码（解决Windows控制台编码问题）
import locale
try:
    # 尝试设置控制台编码
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    if hasattr(sys.stderr, 'reconfigure'):
        sys.stderr.reconfigure(encoding='utf-8')
except Exception:
    pass


# 默认格式参数
DEFAULT_FORMAT = {
    'font': '仿宋_GB2312',
    'body-size': 14,      # pt (四号)
    'title-size': 16,     # pt (三号)
    'sub-size': 15,       # pt (小三)
    'line-spacing': 28,   # pt
    'margins': 2.54,      # cm
    'first-line-indent': 2,  # 字符
    'page-width': 21,     # cm
    'page-height': 29.7,  # cm
}


def parse_format_block(content):
    """从Markdown内容中解析格式指令"""
    fmt = DEFAULT_FORMAT.copy()

    # 匹配 <!-- doc-format ... -->
    pattern = r'<!--\s*doc-format\s*([\s\S]*?)-->'
    match = re.search(pattern, content)
    if not match:
        return fmt, content

    block = match.group(1)
    content = re.sub(pattern, '', content, count=1).strip()

    # 解析各行
    for line in block.split('\n'):
        line = line.strip()
        if ':' not in line:
            continue
        key, value = line.split(':', 1)
        key = key.strip().lower()
        value = value.strip()

        if key == 'font':
            fmt['font'] = value
        elif key == 'body-size':
            v = parse_pt(value)
            if v:
                fmt['body-size'] = v
        elif key == 'title-size':
            v = parse_pt(value)
            if v:
                fmt['title-size'] = v
        elif key == 'sub-size':
            v = parse_pt(value)
            if v:
                fmt['sub-size'] = v
        elif key == 'line-spacing':
            v = parse_pt(value)
            if v:
                fmt['line-spacing'] = v
        elif key == 'margins':
            v = parse_cm(value)
            if v:
                fmt['margins'] = v
        elif key == 'first-line-indent':
            v = parse_pt(value)
            if v:
                fmt['first-line-indent'] = v / 12  # 转换为字符数

    return fmt, content


def parse_pt(value):
    """解析pt值"""
    s = str(value).strip()
    if s.endswith('pt'):
        s = s[:-2]
    try:
        return float(s)
    except:
        return None


def parse_cm(value):
    """解析cm值"""
    s = str(value).strip()
    if s.endswith('cm'):
        s = s[:-2]
    try:
        return float(s)
    except:
        return None


def set_font(run, font_name, font_size, bold=False):
    """设置字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)

    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def set_cell_shading(cell, fill_color="FFFFFF"):
    """设置单元格背景色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), fill_color)
    tcPr.append(shading)


def set_table_border(table):
    """设置表格边框"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def parse_markdown_table(line, lines, start_idx):
    """解析Markdown表格"""
    headers = [h.strip() for h in line.split('|') if h.strip()]

    align_line = lines[start_idx + 1] if start_idx + 1 < len(lines) else ""
    alignments = []
    for cell in align_line.split('|'):
        cell = cell.strip()
        if cell.startswith(':') and cell.endswith(':'):
            alignments.append('center')
        elif cell.endswith(':'):
            alignments.append('right')
        else:
            alignments.append('left')

    data_start = start_idx + 2
    rows_data = []
    for i in range(data_start, len(lines)):
        l = lines[i].strip()
        if not l or not l.startswith('|'):
            return headers, alignments, rows_data, i
        row = [c.strip() for c in l.split('|') if c.strip()]
        rows_data.append(row)

    return headers, alignments, rows_data, len(lines)


def add_table(doc, headers, alignments, rows, fmt):
    """添加表格"""
    if not headers:
        return None

    col_count = len(headers)
    table = doc.add_table(rows=len(rows) + 1, cols=col_count)
    set_table_border(table)

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        para = cell.paragraphs[0]
        run = para.add_run(h)
        set_font(run, '黑体', 12, bold=True)
        set_cell_shading(cell, "D9E2F3")

        align = alignments[i] if i < len(alignments) else 'center'
        if align == 'center':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 数据行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            if c_idx >= col_count:
                break
            cell = table.rows[r_idx + 1].cells[c_idx]
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            set_font(run, fmt['font'], 10.5)

            align = alignments[c_idx] if c_idx < len(alignments) else 'center'
            if align == 'center':
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    return table


def set_heading_style(para, level):
    """设置标题大纲级别"""
    pPr = para._element.get_or_add_pPr()
    outlineLvl = OxmlElement('w:outlineLvl')
    outlineLvl.set(qn('w:val'), str(level - 1))
    pPr.append(outlineLvl)


def convert_md_to_word(input_md_path, output_docx_path, fmt=None):
    """将Markdown文件转换为Word文档"""

    if not os.path.exists(input_md_path):
        print(f'❌ 错误：输入文件不存在：{input_md_path}')
        return False

    with open(input_md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # 解析格式指令
    if fmt is None:
        fmt, content = parse_format_block(content)

    print(f'📋 格式参数：字体={fmt["font"]} | 正文={fmt["body-size"]}pt | 行距={fmt["line-spacing"]}pt | 边距={fmt["margins"]}cm')

    doc = Document()

    # 设置页面
    for section in doc.sections:
        section.top_margin = Cm(fmt['margins'])
        section.bottom_margin = Cm(fmt['margins'])
        section.left_margin = Cm(fmt['margins'])
        section.right_margin = Cm(fmt['margins'])
        section.page_height = Cm(fmt['page-height'])
        section.page_width = Cm(fmt['page-width'])

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # 处理代码块
        if line.strip().startswith('```'):
            if in_code_block:
                para = doc.add_paragraph()
                run = para.add_run('\n'.join(code_lines))
                set_font(run, 'Courier New', 10)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 处理表格
        if line.strip().startswith('|') and '---' not in line:
            headers, alignments, rows_data, end_idx = parse_markdown_table(line, lines, i)
            if headers and rows_data:
                add_table(doc, headers, alignments, rows_data, fmt)
                i = end_idx
                continue

        # 处理占位符
        if line.strip().startswith('[') and '占位符' in line:
            # 添加占位符段落（带特殊格式）
            para = doc.add_paragraph()
            run = para.add_run(line)
            set_font(run, '黑体', 12, bold=True)
            run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
            i += 1
            continue

        # 处理标题
        if line.startswith('# '):
            para = doc.add_paragraph()
            run = para.add_run(line[2:])
            set_font(run, '黑体', fmt['title-size'], bold=True)
            set_heading_style(para, 1)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            para = doc.add_paragraph()
            run = para.add_run(line[3:])
            set_font(run, '黑体', fmt['sub-size'], bold=True)
            set_heading_style(para, 2)
        elif line.startswith('### '):
            para = doc.add_paragraph()
            run = para.add_run(line[4:])
            set_font(run, '黑体', fmt['body-size'], bold=True)
            set_heading_style(para, 3)
        elif line.startswith('#### '):
            para = doc.add_paragraph()
            run = para.add_run(line[5:])
            set_font(run, fmt['font'], fmt['body-size'], bold=True)
            set_heading_style(para, 4)
        elif line.strip() == '':
            i += 1
            continue
        elif re.match(r'^[-*_]{3,}$', line.strip()):
            # 跳过水平线
            i += 1
            continue
        else:
            # 正文段落
            para = doc.add_paragraph()
            run = para.add_run(line)
            set_font(run, fmt['font'], fmt['body-size'])
            para.paragraph_format.first_line_indent = Pt(fmt['body-size'] * fmt['first-line-indent'])
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        i += 1

    # 设置行距
    for para in doc.paragraphs:
        para.paragraph_format.line_spacing = Pt(fmt['line-spacing'])
        para.paragraph_format.line_spacing_rule = 3
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

    doc.save(output_docx_path)
    print(f'✅ Word文档已生成：{output_docx_path}')
    return True


def main():
    if len(sys.argv) != 3:
        print('用法: python3 convert_to_docx.py <输入.md> <输出.docx>')
        print('示例: python3 convert_to_docx.py tech_spec.md output.docx')
        sys.exit(1)

    success = convert_md_to_word(sys.argv[1], sys.argv[2])
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
