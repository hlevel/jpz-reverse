#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOC/DOCX转Markdown转换脚本
将Word文档转换为Markdown格式，用于标书知识库构建
版本：v1.0 (2026-05-21)
"""

import io
import os
import sys
from pathlib import Path

# 设置标准输出编码（解决Windows控制台编码问题）
import locale
try:
    # 尝试设置控制台编码
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    if hasattr(sys.stderr, 'reconfigure'):
        sys.stderr.reconfigure(encoding='utf-8')
except Exception:
    pass


def extract_docx_content(file_path: str) -> str:
    """从docx文件提取内容"""
    try:
        from docx import Document
        doc = Document(file_path)
        content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # 检测标题样式
                style_name = para.style.name.lower()
                if 'heading 1' in style_name:
                    content.append(f'# {text}')
                elif 'heading 2' in style_name:
                    content.append(f'## {text}')
                elif 'heading 3' in style_name:
                    content.append(f'### {text}')
                elif 'heading 4' in style_name:
                    content.append(f'#### {text}')
                else:
                    content.append(text)

        # 提取表格
        for i, table in enumerate(doc.tables, 1):
            content.append(f'\n### 表格 {i}\n')
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                content.append(f'| {row_text} |')

        return '\n\n'.join(content)
    except Exception as e:
        return f"[错误] 无法读取 {file_path}: {str(e)}"


def extract_doc_content(file_path: str) -> str:
    """从doc文件提取内容（需要pywin32）"""
    try:
        import win32com.client
        import pythoncom

        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        try:
            doc = word.Documents.Open(os.path.abspath(file_path))
            content = doc.Content.Text
            doc.Close(False)
            return content
        finally:
            word.Quit()
            pythoncom.CoUninitialize()
    except ImportError:
        return "[跳过] pywin32未安装，无法读取.doc文件"
    except Exception as e:
        return f"[错误] 无法读取 {file_path}: {str(e)}"


def process_doc_folder(source_dir: str, output_dir: str, prefix: str = "knowledge"):
    """处理DOC/DOCX文件夹"""
    os.makedirs(output_dir, exist_ok=True)

    # 获取所有DOC/DOCX文件
    doc_files = []
    for root, dirs, files in os.walk(source_dir):
        for file in files:
            if file.lower().endswith(('.doc', '.docx')):
                doc_files.append(os.path.join(root, file))

    if not doc_files:
        print(f"[WARNING] 在 {source_dir} 中没有找到DOC/DOCX文件")
        return

    print(f"找到 {len(doc_files)} 个DOC/DOCX文件")
    print("=" * 60)

    processed_count = 0
    for doc_path in doc_files:
        doc_name = os.path.basename(doc_path)
        print(f"\n处理: {doc_name}")

        # 提取内容
        if doc_name.lower().endswith('.docx'):
            content = extract_docx_content(doc_path)
        else:
            content = extract_doc_content(doc_path)

        # 生成输出文件名
        md_name = doc_name.rsplit('.', 1)[0] + '.md'
        md_path = os.path.join(output_dir, f"{prefix}_{md_name}")

        # 保存为md文件
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(f"# {doc_name}\n\n")
            f.write(f"**来源**: {doc_path}\n\n")
            f.write("---\n\n")
            f.write(content)

        print(f"  已保存: {os.path.basename(md_path)}")
        processed_count += 1

    print("\n" + "=" * 60)
    print(f"处理完成！成功转换 {processed_count}/{len(doc_files)} 个DOC/DOCX文件")
    print("=" * 60)


def main():
    import argparse

    parser = argparse.ArgumentParser(description='DOC/DOCX转Markdown转换工具')
    parser.add_argument('--source', '-s',
                       required=True,
                       help='源DOC/DOCX文件夹路径')
    parser.add_argument('--output', '-o',
                       default='./templates',
                       help='输出文件夹路径')
    parser.add_argument('--prefix', '-p',
                       default='knowledge',
                       help='输出文件名前缀 (默认: knowledge)')

    args = parser.parse_args()

    if not os.path.exists(args.source):
        print(f"[ERROR] 源文件夹不存在: {args.source}")
        sys.exit(1)

    process_doc_folder(args.source, args.output, args.prefix)


if __name__ == '__main__':
    main()